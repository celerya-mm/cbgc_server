"""
	In ambiente Linux per funzionare c'è bisogno che queste utilities siano installate:
	apt-get install -y apt-utils --no-install-recommends
	apt-get install -y dialog
	apt-get install -y libreoffice --no-install-recommends
	apt-get install python3-uno
	apt-get install -y default-jre
	apt-get install -y unoconv
"""

import base64
import os
import platform
import shutil
import subprocess as subp
from subprocess import call

import qrcode
from docx import Document  # noqa
from docx.shared import Inches  # noqa
from docxtpl import DocxTemplate

# imposta path file's models docx (Windows)
path_docx = os.path.dirname(os.path.realpath(__file__))
folders_models_docx = os.path.join(path_docx, "certificates_models_docx")

# imposta path file's models .odt (Linux)
path_odt = os.path.dirname(os.path.realpath(__file__))
folders_models_odt = os.path.join(path_odt, "certificates_models_odt")

# imposta path work
folder_work = os.path.join(path_docx, "certificates_work")
if not os.path.exists(folder_work):
	os.makedirs(folder_work)

# imposta path temp file
folder_temp = os.path.join(path_docx, "temp_pdf")
if not os.path.exists(folder_temp):
	os.makedirs(folder_temp)


def generate_qr_code(_str):
	"""Genera un QR-Code da una stringa."""
	try:
		qr = qrcode.QRCode(
			version=1,
			error_correction=qrcode.constants.ERROR_CORRECT_H,
			box_size=10,
			border=1,
		)
		qr.add_data(_str.replace("/", "_"))
		qr.make(fit=True)
		img = qr.make_image(fill_color="black", back_color="white").convert('RGB')
		# print("")
		# print(img)
		img.save(os.path.join(folder_work, _str.replace("/", "_") + ".jpg"))
		# print("QR-Code Generated")
		return _str
	except Exception as err:
		print(err)
		return False


def byte_to_pdf(byte, f_name):
	"""Ricrea il pdf da una stringa in byte."""
	# svuoto la cartella da file vecchio
	for filename in os.listdir(folder_temp):
		file_path = os.path.join(folder_temp, filename)
		os.remove(file_path)

	path_file = os.path.join(folder_temp, f_name.replace("/", "_") + ".pdf")
	# print("PDF_STR_TYPE:", type(byte), "LEN:", len(byte))
	with open(path_file, "wb") as f:
		f.write(byte)
	return path_file


def docx_to_pdf(target):
	"""Converte file Word in pdf."""
	try:
		from comtypes import client
		word = client.CreateObject('Word.Application')
	except ImportError as err:
		print(err)
		return False

	doc = word.Documents.Open(target)
	outFile = target.replace("docx", "pdf")
	# print("PDF_PATH:", outFile)
	# print("")
	try:
		doc.ExportAsFixedFormat(
			OutputFileName=outFile,
			ExportFormat=17,  # 17 = PDF output, 18=XPS output
			OpenAfterExport=False,
			OptimizeFor=0,  # 0=Print (higher res), 1=Screen (lower res)
			CreateBookmarks=1,  # 0=No bookmarks, 1=Heading bookmarks only, 2=bookmarks match word bookmarks
			DocStructureTags=False
		)
		doc.Close()
		os.remove(target)
		print("SUCCESSFULLY_CONVERTED")
		return outFile
	except Exception as err:
		os.remove(target)
		print("ERROR_CONVERT:", err)
		return False


def odt_to_pdf(target):
	"""Converte file LibreOffice in pdf."""
	try:
		outPath = target.replace("odt", "pdf").split("/")
		outPath = os.path.join(folder_temp, outPath[len(outPath) - 1])
		print("OUT_NAME:", outPath)
		call(["libreoffice", "--headless", "--convert-to", "pdf", target, "--outdir", folder_temp])
		print("SUCCESSFULLY_CONVERTED")
		os.remove(target)
		print("OUT_FILE:", outPath)
		return outPath
	except Exception as err:
		print("ERROR_CONVERT:", err)
		return False


def insert_qrcode(target, _str):
	"""Inserisce QR-Code dopo marker in file."""
	try:
		_jpg = os.path.join(folder_work, _str.replace("/", "_") + ".jpg")
		document = Document(target)
		for para in document.paragraphs:
			if '{{QRCode}}' in para.text:
				# print("PARA_TEXT:", para.text)
				# create a new run (a new line of text)
				run = para.add_run()
				# insert the image
				run.add_picture(_jpg, width=Inches(1.1), height=Inches(1.1))
				os.remove(_jpg)
		# save the modified docx file
		document.save(target)
		return True
	except Exception as err:
		print("ERRORE_INSERIMENTO_QRCOD:", err)
		return False


def insert_qrcode_linux(target, _str):
	"""Inserisce QR-Code dopo marker in file odt."""
	try:
		_jpg = os.path.join(folder_work, _str.replace("/", "_") + ".jpg")
		# print("QRCODE_FIND")

		# Read the file into a string variable
		with open(target, "r", encoding="ISO-8859-1") as model:
			content = model.read()
		# print("CONTENT_READ")

		# Replace the marker with the image content
		img = f'<p><img src={_jpg} alt="QRCode"></p>'
		content = content.replace("{{QRCode}}", img)
		# print("CONTENT_REPLACED")

		# Write the updated content back to the file
		with open(target, "w") as model:
			model.write(content)
		# print("SAVED")
		return True
	except Exception as err:
		print("ERRORE_INSERIMENTO_QRCOD:", err)
		return False


def pdf_to_byte(_pdf):
	"""Converte pdf in byte string."""
	with open(_pdf, "rb") as f:
		b_string = str(base64.b64encode(f.read()), 'utf-8')
		# print(enc_string[100], type(enc_string))
		# converto in string utf-8 per caricare su db
		b_string = base64.b64decode(b_string)
	os.remove(_pdf)
	return b_string


def insert_data(target, data):
	"""Inserisce dati in file sostituendo i markers."""
	try:
		print("TARGET:", target, "data:", data)
		tpl = DocxTemplate(target)  # leggi il template word
		tpl.render(data)
		# save the modified docx file
		tpl.save(target)
		return True
	except Exception as err:
		print("ERRORE_INSERMENTO_DATI_IN_FILE:", err)
		return False


def insert_data_linux(target, data):
	"""Inserisce dati in file sostituendo i markers."""
	try:
		# print("TARGET:", target, "data:", data)
		# Convert the ODT file to plain text
		p = subp.run(["unoconv", "-f", "txt", target], capture_output=True, text=True)
		text = p.stdout
		print("READED:", str(text))
		# Loop through the paragraphs and replace markers
		for marker, replacement in data.items():
			text = text.replace("{%s}" % str(marker), str(replacement))
		print("REPLACED:", text)
		# Save the modified plain text to a temporary file
		with open("modified.txt", "w") as f:
			f.write(text)
		# Convert the modified plain text back to ODT format
		subp.run(["unoconv", "-f", "odt", "-o", target, "modified.txt"])
		print("SUBSTITUTED:", target)
		return True
	except Exception as err:
		print("ERRORE_INSERMENTO_DATI_IN_FILE:", err)
		return False


def create_b_certificate_docx(data, _file, _str):
	"""Legge il modello dal file word e sostituisce le chiavi."""
	# copio il modello e lo salvo nella cartella di lavoro
	_name = data["certificate_nr"].replace("/", "_") + ".docx"
	target = os.path.join(folder_work, _name)
	source = os.path.join(folders_models_docx, _file)
	shutil.copy(source, target)
	# inserisce QR_code
	if insert_qrcode(target, _str):
		# inserisce dati nel modello
		if insert_data(target, data):
			# convert docx to pdf
			_pdf = docx_to_pdf(target)
			# convert pdf to byte
			if _pdf is not False:
				return pdf_to_byte(_pdf)
	else:
		return False


def create_b_certificate_odt(data, _file, _str):
	"""Legge il modello dal file LibreOffice e sostituisce le chiavi."""
	# copio il modello e lo salvo nella cartella di lavoro
	_name = data["certificate_nr"].replace("/", "_") + ".odt"
	target = os.path.join(folder_work, _name)
	source = os.path.join(folders_models_odt, _file)
	shutil.copy(source, target)
	# inserisce QR_code
	try:
		if insert_qrcode_linux(target, _str):
			# inserisce dati nel modello
			if insert_data_linux(target, data):
				# convert odt to pdf
				_pdf = odt_to_pdf(target)
				# convert pdf to byte
				if _pdf is not False:
					return pdf_to_byte(_pdf)
		else:
			return False
	except Exception as err:
		print(err)
		return False


def create_pdf_certificate(buyer_type, _data, str_qr):
	if buyer_type == "Ristorante":
		if "Windows" in platform.platform():
			model_name = "certificato_ristoranti.docx"
			pdf_str = create_b_certificate_docx(_data, model_name, str_qr)
		elif "Linux" in platform.platform():
			model_name = "certificato_ristoranti.odt"
			pdf_str = create_b_certificate_odt(_data, model_name, str_qr)
		else:
			print("PIATTAFORMA_NON_RICONOSCIUTA")
			return False
	else:
		if "Windows" in platform.platform():
			model_name = "certificato_macellerie.docx"
			pdf_str = create_b_certificate_docx(_data, model_name, str_qr)
		elif "Linux" in platform.platform():
			print("PLATFORM:", str(platform.platform()))
			model_name = "certificato_macellerie.odt"
			pdf_str = create_b_certificate_odt(_data, model_name, str_qr)
		else:
			print("PIATTAFORMA_NON_RICONOSCIUTA")
			return False

	return pdf_str
